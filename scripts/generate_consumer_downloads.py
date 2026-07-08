from pathlib import Path
import json, re, shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT/'public'/'downloads'
DOWNLOADS.mkdir(parents=True, exist_ok=True)

# Remove consumer-hostile public markdown/text outputs. Keep source markdown under templates/prompts only.
for p in DOWNLOADS.glob('*.md'):
    p.unlink()
for p in DOWNLOADS.glob('*.txt'):
    p.unlink()

catalog = json.loads((ROOT/'data/products/product_catalog.json').read_text())['products']
products = [p for p in catalog if p.get('status') == 'active_paid']

def product_by_sku(sku):
    return next(p for p in products if p['sku']==sku)

SPEC = {
'letter-of-explanation': {
  'use': 'Use this when someone asks you to explain a gap, issue, change, late payment, address history, credit note, income gap, or other application question.',
  'letter_title': 'Letter of Explanation Template',
  'letter_intro': 'I am writing to explain [short issue] and to share the steps I am taking now.',
  'facts': ['What happened?', 'When did it happen?', 'What changed?', 'What proof can you attach?', 'What are you doing next?'],
  'attach': ['Any notice or request you received', 'Proof that supports your explanation', 'Updated income, address, or payment records if helpful'],
  'steps': ['Read the request from the landlord, lender, school, agency, or company.', 'Write only true facts.', 'Keep the letter short.', 'Attach only real documents.', 'Save a copy before you send it.']
},
'income-employment-letter-kit': {
  'use': 'Use this when you need to explain work, income, a job change, a gap in work, self-employment, or bank deposits.',
  'letter_title': 'Income and Employment Explanation Template',
  'letter_intro': 'I am writing to explain my current work and income information.',
  'facts': ['Who pays you?', 'How often are you paid?', 'What changed recently?', 'Do you have a work gap?', 'What proof shows your income?'],
  'attach': ['Pay stubs', 'Offer letter or employment letter', 'Bank statements only if requested', 'Tax documents or invoices if self-employed'],
  'steps': ['Choose the income letter that fits your situation.', 'Use real dates and amounts.', 'Do not round up income.', 'Attach matching proof.', 'Ask the recipient where to send it.']
},
'credit-letter-kit': {
  'use': 'Use this when you want to prepare your own credit letters, dispute possible errors, ask for goodwill, request debt validation, or explain credit history.',
  'letter_title': 'Self-Service Credit Letter Template',
  'letter_intro': 'I am writing about an item on my credit report or credit file.',
  'facts': ['What item are you writing about?', 'Why do you believe it is wrong or needs review?', 'What proof do you have?', 'What exact action are you asking for?', 'Where should the reply be sent?'],
  'attach': ['Credit report page showing the item', 'Proof of payment or identity when needed', 'Any creditor letter or account notice'],
  'steps': ['Get your credit report.', 'Circle or write down the exact item.', 'Do not dispute accurate facts just because they hurt.', 'Attach proof.', 'Mail or submit the letter yourself and keep copies.']
},
'rental-application-kit': {
  'use': 'Use this before you apply for an apartment or rental home and need your papers organized.',
  'letter_title': 'Rental Application Cover Letter Template',
  'letter_intro': 'I am applying for the rental home and have included documents to support my application.',
  'facts': ['Who will live there?', 'What income proof will you send?', 'Do you need to explain rental history?', 'Do you need to explain credit or income?', 'What documents did the landlord ask for?'],
  'attach': ['ID if requested', 'Pay stubs or income proof', 'Rental history details', 'References if requested', 'Explanation letter if needed'],
  'steps': ['Read the rental listing and application rules.', 'Make a document folder.', 'Prepare your cover letter.', 'Attach real proof.', 'Send the packet yourself.']
},
'loan-prep-letter-kit': {
  'use': 'Use this before an auto loan, personal loan, mortgage prep, or financing review when you need to explain documents.',
  'letter_title': 'Loan Prep Explanation Template',
  'letter_intro': 'I am writing to explain the documents connected to my financing review.',
  'facts': ['What loan or financing request is this for?', 'What question did the lender ask?', 'What income, credit, or bank activity needs context?', 'What proof can you attach?', 'What changed or improved?'],
  'attach': ['Income proof', 'Bank statements if requested', 'Credit explanation if needed', 'Employment proof', 'Any lender request letter'],
  'steps': ['List the lender questions.', 'Answer only what was asked.', 'Use real numbers.', 'Attach proof that matches the answer.', 'Keep a copy of every file sent.']
},
'business-funding-prep-kit': {
  'use': 'Use this before business funding, SBA prep, or lender review when you need to organize business documents and explain business facts.',
  'letter_title': 'Business Funding Prep Letter Template',
  'letter_intro': 'I am writing to organize and explain my business information for review.',
  'facts': ['What does the business do?', 'How does the business make money?', 'What funding amount or use of funds are you explaining?', 'What documents show revenue or cash flow?', 'What questions did the lender ask?'],
  'attach': ['Business bank statements if requested', 'Tax documents if requested', 'Invoices or revenue proof', 'Business registration documents', 'Use-of-funds notes'],
  'steps': ['Make a list of required funding documents.', 'Explain the business in simple words.', 'Use real revenue and expense numbers.', 'Attach only documents you are allowed to share.', 'Save a clean copy of the full packet.']
},
'life-admin-letter-kit': {
  'use': 'Use this for everyday paperwork, address history, residency, identity details, school requests, employer requests, account issues, or records requests.',
  'letter_title': 'Life Admin Letter Template',
  'letter_intro': 'I am writing to explain or request help with a paperwork matter.',
  'facts': ['What office, school, company, or agency is this for?', 'What do they need from you?', 'What facts do you need to explain?', 'What proof do you have?', 'What do you want them to do next?'],
  'attach': ['Proof of address', 'ID if requested', 'Account notice or letter', 'School or employer form', 'Any real document that supports the request'],
  'steps': ['Write the request in one sentence.', 'Add true details.', 'Attach only useful proof.', 'Ask for a clear next step.', 'Save the letter and response.']
},
'complete-approvalprep-bundle': {
  'use': 'Use this when you want the full ApprovalPrep library for rental, credit, income, loan, business funding, and life admin paperwork.',
  'letter_title': 'Complete ApprovalPrep Bundle - Start Here',
  'letter_intro': 'This bundle helps you choose the right letter or checklist for your situation.',
  'facts': ['What are you applying for or responding to?', 'Which person or office asked for documents?', 'Which kit fits the situation?', 'What facts are true and provable?', 'What deadline do you have?'],
  'attach': ['Only documents requested or clearly helpful', 'Proof that matches the letter', 'A copy of the request or notice', 'Your completed checklist'],
  'steps': ['Choose the kit that matches your situation.', 'Start with the checklist.', 'Fill in one letter at a time.', 'Review the safety page.', 'Send the final packet yourself.']
}
}

DISCLAIMER_TITLE = 'Important: Read This First'
DISCLAIMER = [
'ApprovalPrep is a self-service paperwork tool. You fill in your own information. You choose what to send. You send it yourself.',
'ApprovalPrep is not a law firm, lender, broker, credit repair company, accountant, employment verification company, rental agent, or financial advisor.',
'ApprovalPrep does not contact landlords, lenders, employers, credit bureaus, creditors, agencies, schools, or companies for you.',
'ApprovalPrep does not create fake documents, change documents, verify facts for you, remove accurate credit items, or promise approval, funding, housing, credit results, or any other outcome.',
'Use only true facts and real documents. Do not send anything false, edited to mislead, or not yours to send.',
'This kit is educational and organizational only. Ask a licensed professional if you need legal, tax, accounting, credit, lending, housing, immigration, employment, or financial advice.'
]

COMMON_SECTIONS = ['Start Here', 'Before You Fill This In', 'Step-by-Step Instructions', 'Fill-In Template', 'Self-Review Checklist', 'What to Attach', 'What to Do Next', 'What ApprovalPrep Does Not Do']

def add_docx_disclaimer(doc, product):
    sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
    h = doc.add_heading(product['name'], 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Plain-language self-service paperwork kit')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(DISCLAIMER_TITLE, level=1)
    for d in DISCLAIMER:
        para = doc.add_paragraph(d)
        para.style = 'List Bullet'
    doc.add_paragraph('By using this kit, you agree that you are responsible for your own documents, choices, and results.')
    doc.add_page_break()

def add_docx_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Question'
    hdr[1].text = 'Your answer'
    for q in rows:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = ''
    return table

def make_docx(product):
    spec = SPEC[product['sku']]
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(11)
    styles['Heading 1'].font.name = 'Arial'; styles['Heading 1'].font.size = Pt(18)
    styles['Heading 2'].font.name = 'Arial'; styles['Heading 2'].font.size = Pt(14)
    add_docx_disclaimer(doc, product)
    doc.add_heading('Start Here', level=1)
    doc.add_paragraph(spec['use'])
    doc.add_paragraph('This kit gives you a clear path. It does not make a decision for the person reading your papers.')
    doc.add_heading('Before You Fill This In', level=1)
    for item in ['Use simple words.', 'Use true facts only.', 'Do not guess.', 'Do not hide important facts.', 'Do not send private documents unless the recipient asked for them or they clearly support your letter.']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('Step-by-Step Instructions', level=1)
    for i, item in enumerate(spec['steps'], 1):
        doc.add_paragraph(f'Step {i}: {item}')
    doc.add_heading('Fill-In Template', level=1)
    doc.add_heading(spec['letter_title'], level=2)
    doc.add_paragraph('Date: ____________________')
    doc.add_paragraph('To: ____________________')
    doc.add_paragraph('Re: ____________________')
    doc.add_paragraph('Hello,')
    doc.add_paragraph(spec['letter_intro'])
    doc.add_paragraph('Here are the facts:')
    for q in spec['facts']:
        doc.add_paragraph(q + ' ________________________________________________')
    doc.add_paragraph('The documents I am including are: ________________________________________________')
    doc.add_paragraph('Thank you for reviewing this information. Please let me know if you need anything else from me.')
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('Name: ____________________')
    doc.add_heading('Self-Review Checklist', level=1)
    for item in ['I used my real name and correct contact information.', 'Every date and dollar amount is correct.', 'I removed anything that is not true.', 'I attached only real documents.', 'I saved a copy for myself.']:
        doc.add_paragraph('☐ ' + item)
    doc.add_heading('What to Attach', level=1)
    for item in spec['attach']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('Questions to Answer Before You Send', level=1)
    add_docx_table(doc, spec['facts'])
    doc.add_heading('Do and Do Not', level=1)
    for item in ['Do keep the letter short.', 'Do answer the question that was asked.', 'Do attach proof when it helps.', 'Do not make up income, dates, addresses, payments, or documents.', 'Do not promise that someone will approve you.']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('Follow-Up Message Template', level=1)
    doc.add_paragraph('Hello, I am following up on the documents I sent on [date]. Please let me know if anything else is needed from me. Thank you.')
    doc.add_heading('What to Do Next', level=1)
    for item in ['Send the letter yourself.', 'Write down the date you sent it.', 'Save the receipt, email, upload confirmation, or tracking number.', 'Follow up politely if you do not hear back.', 'Keep all replies in one folder.']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('What ApprovalPrep Does Not Do', level=1)
    for item in DISCLAIMER[1:6]:
        doc.add_paragraph(item, style='List Bullet')
    out = DOWNLOADS/f"{product['sku']}.docx"
    doc.save(out)
    return out

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='Small', fontSize=8.8, leading=11, spaceAfter=5))
styles.add(ParagraphStyle(name='H1x', parent=styles['Heading1'], fontSize=17, leading=20, spaceAfter=8))
styles.add(ParagraphStyle(name='H2x', parent=styles['Heading2'], fontSize=13, leading=16, spaceAfter=6))
styles.add(ParagraphStyle(name='Bodyx', parent=styles['BodyText'], fontSize=10.5, leading=14, spaceAfter=6))

def P(text, style='Bodyx'):
    return Paragraph(text.replace('&','&amp;'), styles[style])

def bullets(items):
    flow=[]
    for item in items:
        flow.append(P('• ' + item, 'Bodyx'))
    return flow

def make_pdf(product):
    spec = SPEC[product['sku']]
    out = DOWNLOADS/f"{product['sku']}.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=letter, topMargin=.55*inch, bottomMargin=.55*inch, leftMargin=.65*inch, rightMargin=.65*inch)
    flow = []
    flow.append(P(product['name'], 'H1x'))
    flow.append(P('Plain-language self-service paperwork kit', 'Bodyx'))
    flow.append(P(DISCLAIMER_TITLE, 'H2x'))
    flow.extend(bullets(DISCLAIMER))
    flow.append(P('By using this kit, you agree that you are responsible for your own documents, choices, and results.', 'Bodyx'))
    flow.append(PageBreak())
    flow.append(P('Start Here', 'H1x'))
    flow.append(P(spec['use']))
    flow.append(P('This kit gives you a clear path. It does not make a decision for the person reading your papers.'))
    flow.append(P('Before You Fill This In', 'H1x'))
    flow.extend(bullets(['Use simple words.', 'Use true facts only.', 'Do not guess.', 'Do not hide important facts.', 'Do not send private documents unless the recipient asked for them or they clearly support your letter.']))
    flow.append(P('Step-by-Step Instructions', 'H1x'))
    for i,item in enumerate(spec['steps'],1): flow.append(P(f'Step {i}: {item}'))
    flow.append(P('Fill-In Template', 'H1x'))
    flow.append(P(spec['letter_title'], 'H2x'))
    for line in ['Date: ____________________', 'To: ____________________', 'Re: ____________________', 'Hello,', spec['letter_intro'], 'Here are the facts:']:
        flow.append(P(line))
    for q in spec['facts']:
        flow.append(P(q + ' ________________________________________________'))
    for line in ['The documents I am including are: ________________________________________________', 'Thank you for reviewing this information. Please let me know if you need anything else from me.', 'Sincerely,', 'Name: ____________________']:
        flow.append(P(line))
    flow.append(P('Self-Review Checklist', 'H1x'))
    flow.extend(bullets(['I used my real name and correct contact information.', 'Every date and dollar amount is correct.', 'I removed anything that is not true.', 'I attached only real documents.', 'I saved a copy for myself.']))
    flow.append(P('What to Attach', 'H1x'))
    flow.extend(bullets(spec['attach']))
    flow.append(P('Questions to Answer Before You Send', 'H1x'))
    table_data = [[P('Question','Small'), P('Your answer','Small')]] + [[P(q,'Small'), P('','Small')] for q in spec['facts']]
    tbl = Table(table_data, colWidths=[2.6*inch, 3.6*inch], rowHeights=[.3*inch]+[.5*inch]*len(spec['facts']))
    tbl.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.5,colors.grey),('VALIGN',(0,0),(-1,-1),'TOP'),('BACKGROUND',(0,0),(-1,0),colors.whitesmoke)]))
    flow.append(tbl)
    flow.append(Spacer(1, .15*inch))
    flow.append(P('Do and Do Not', 'H1x'))
    flow.extend(bullets(['Do keep the letter short.', 'Do answer the question that was asked.', 'Do attach proof when it helps.', 'Do not make up income, dates, addresses, payments, or documents.', 'Do not promise that someone will approve you.']))
    flow.append(P('Follow-Up Message Template', 'H1x'))
    flow.append(P('Hello, I am following up on the documents I sent on [date]. Please let me know if anything else is needed from me. Thank you.'))
    flow.append(P('What to Do Next', 'H1x'))
    flow.extend(bullets(['Send the letter yourself.', 'Write down the date you sent it.', 'Save the receipt, email, upload confirmation, or tracking number.', 'Follow up politely if you do not hear back.', 'Keep all replies in one folder.']))
    flow.append(P('What ApprovalPrep Does Not Do', 'H1x'))
    flow.extend(bullets(DISCLAIMER[1:6]))
    doc.build(flow)
    return out

made=[]
for p in products:
    made.append(str(make_docx(p)))
    made.append(str(make_pdf(p)))

# Update function to expose only consumer friendly files.
verify_path = ROOT/'functions/api/verify-download.js'
verify = verify_path.read_text()
verify = re.sub(r'files:\s*\[\s*`/downloads/\$\{sku\}\.pdf`,\s*`/downloads/\$\{sku\}\.docx`,\s*`/downloads/\$\{sku\}-guide\.md`,\s*`/downloads/\$\{sku\}-worksheet\.txt`\s*\]', 'files: [\n      `/downloads/${sku}.pdf`,\n      `/downloads/${sku}.docx`\n    ]', verify, flags=re.S)
verify_path.write_text(verify)

# Update download page copy.
page = ROOT/'src/pages/[...slug].astro'
text = page.read_text()
text = text.replace('Production downloads require Stripe payment verification before the PDF, DOCX, guide, and worksheet links are released.', 'Production downloads require Stripe payment verification before the PDF and editable DOCX files are released.')
page.write_text(text)

# Update button labels for download names.
layout = ROOT/'src/layouts/BaseLayout.astro'
text = layout.read_text()
text = text.replace('a.textContent = file.split("/").pop();', 'a.textContent = file.endsWith(".pdf") ? "Download PDF guide" : "Download editable Word document";')
layout.write_text(text)

# Update manifest for 8 active product downloads only.
manifest = {'schemaVersion':'3.0.0','delivery':'verified_stripe_download','consumerFormats':['pdf','docx'],'removedConsumerFormats':['markdown','txt'],'reason':'Public paid downloads must be useful to average consumers. Markdown and TXT are source/internal formats only.','products':[]}
for p in products:
    manifest['products'].append({'sku':p['sku'],'name':p['name'],'files':[f'/downloads/{p["sku"]}.pdf',f'/downloads/{p["sku"]}.docx'],'firstPageRequires':['plain-language self-service boundary','not legal/financial/credit-repair/lending advice','no third-party contact','no fake documents','no promises or guarantees'],'requiredSections':COMMON_SECTIONS})
(ROOT/'data/products/download_manifest.json').write_text(json.dumps(manifest, indent=2))

# Add hostile review report.
report = ROOT/'docs'/'DOWNLOAD_PRODUCT_HOSTILE_REVIEW_v4_2_6.md'
report.write_text('''# Download Product Hostile Review v4.2.6\n\nStatus: patched from hostile review.\n\n## Findings\n\n- Markdown guide and TXT worksheet were consumer-hostile formats. Removed from public verified downloads.\n- Prior PDFs were too thin to justify paid value. Rebuilt all 8 PDFs and DOCX files with first-page disclaimers, step-by-step instructions, fill-in templates, checklists, attachment guidance, and next steps.\n- Public downloads now expose PDF and editable Word document only.\n\n## Guardrails now required\n\nEach paid download must use plain language, truthful-document boundaries, no credit repair promises, no legal/financial/lending/accounting advice, no third-party contact, no fake documents, and no guaranteed approval or funding claim.\n\n## Consumer value standard\n\nEach kit must help a normal consumer answer: what is this, when do I use it, what do I fill in, what proof should I attach, what do I send, what do I do next, and what should I never claim.\n''')

print('\n'.join(made))
